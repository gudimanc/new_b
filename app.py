from flask import Flask, render_template, redirect,url_for, request
from flask.helpers import flash
from flask_sqlalchemy import SQLAlchemy
from form import AvailabilityForm, ProductForm,DeleteProduct,EditProduct,EditProductName,BillingForm,ClearBill,UndoBill,GenerateBill
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
import os
import json
import pandas as pd
import time
import datetime
import docx
app = Flask(__name__)
app.config['SECRET_KEY']='LongAndRandomSecretKey'
basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'app.sqlite')


db = SQLAlchemy(app)
migrate = Migrate(app, db)


class Products(db.Model):
    id = db.Column('product_id', db.Integer, primary_key = True)
    name = db.Column(db.String(100))
    price = db.Column(db.Integer())     
    available_quantity = db.Column(db.Integer(), nullable=True)
    def __init__(self,name,price):
        self.name=name
        self.price=price

class Bill(db.Model):
    id = db.Column('bill_id', db.Integer, primary_key = True)
    name = db.Column(db.String(100))
    details = db.Column(db.String(1000))

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/products',methods=['GET','POST'])
def products():
    try:
        add_form = ProductForm()
        delete_form = DeleteProduct()
        edit_form = EditProduct()
        edit_name_form = EditProductName()
        if add_form.validate_on_submit() and add_form.add.data:
            r=Products.query.filter_by(name=add_form.name.data.capitalize()).first()
            if r:
                return redirect(url_for('products'))
            p=Products(add_form.name.data.capitalize(),add_form.price.data)
            db.session.add(p)
            db.session.commit()
            return redirect(url_for('products'))
        if delete_form.validate_on_submit() and delete_form.remove.data:
            r=Products.query.filter_by(name=delete_form.name.data.capitalize()).first()
            db.session.delete(r)
            db.session.commit()
            return redirect(url_for('products'))
        if edit_form.validate_on_submit() and edit_form.edit.data:
            e=Products.query.filter_by(name=edit_form.name.data.capitalize()).first()
            e.price = edit_form.price.data
            db.session.add(e)
            db.session.commit()
            return redirect(url_for('products'))
        if edit_name_form.validate_on_submit() and edit_name_form.edit.data:
            e=Products.query.filter_by(name=edit_name_form.old_name.data.capitalize()).first()
            e.name = edit_name_form.new_name.data.capitalize()
            db.session.add(e)
            db.session.commit()
            return redirect(url_for('products'))
        return render_template('products.html',add_form=add_form,delete_form=delete_form,edit_form=edit_form,edit_name_form=edit_name_form,all_products=Products.query.all())
    except Exception as e:
        print(e)
        return redirect(url_for('products'))

@app.route('/availability',methods=['GET','POST'])
def availability():
    try:
        form = AvailabilityForm()
        all_products = Products.query.all()
        all_names=[]
        for product in all_products:
            all_names.append(product.name)
        form.name.choices=all_names
        if form.validate_on_submit() and form.change.data:
            e=Products.query.filter_by(name=form.name.data).first()
            e.available_quantity = form.available_quantity.data
            db.session.add(e)
            db.session.commit()
            return redirect(url_for('availability'))
        return render_template('availability.html',all_products=all_products,form=form)
    except Exception as e:
        print(e)
        return redirect(url_for('availability'))


@app.route('/billing',methods=['GET','POST'])
def billing():
    try:
        file_name = 'bills/temp.csv'
        amount=0
        form = BillingForm()
        cb =ClearBill()
        ub=UndoBill()
        gb=GenerateBill()
        all_products = Products.query.all()
        all_names=[]
        for product in all_products:
            all_names.append(product.name)
        form.name.choices=all_names
        with open(file_name,'r') as fp:
            outp=fp.readlines()
        outs=[]
        for line in outp:
            outs.append(json.loads(line))
            amount+=json.loads(line)['amount']
        if form.validate_on_submit() and form.add.data:
            e=Products.query.filter_by(name=form.name.data).first()
            temp_json={}
            temp_json['name']=form.name.data
            temp_json['quantity']=form.quantity.data
            temp_json['price']=e.price
            temp_json['amount']=int(e.price)*int(form.quantity.data)
            with open(file_name,'a') as fp:
                fp.write(json.dumps(temp_json)+'\n')
            
            
            return redirect(url_for('billing'))
        if cb.validate_on_submit() and cb.clear.data:
            with open(file_name,'w') as fp:
                pass
            return redirect(url_for('billing'))
        if ub.validate_on_submit() and ub.undo.data:
            with open(file_name,'r') as fp:
                x=fp.readlines()
            print(x)
            x=x[:-1]
            with open(file_name,'w') as fp:
                for i in x:
                    fp.write(i)
            return redirect(url_for('billing'))
        if gb.validate_on_submit() and gb.genbill.data:
            with open(file_name,'r') as fp:
                outp=fp.readlines()
            outs=[]
            o={'name':[],'quantity':[],'price':[],'amount':[],'id':None}
            for line in outp:
                outs.append(json.loads(line))
                o['name'].append(outs[-1]['name'])
                o['quantity'].append(outs[-1]['quantity'])
                o['price'].append(outs[-1]['price'])
                o['amount'].append(outs[-1]['amount'])
            b = Bill()
            b.name = gb.name.data
            b.details = str(o)
            db.session.add(b)
            db.session.commit()
            o['id'] = Bill.query.filter_by(name=gb.name.data,details=str(o)).first().id
            df=pd.DataFrame(o)
            print(df)
            for name in df['name'].values:

                e=Products.query.filter_by(name=name).first()

                orig_quantity = int(e.available_quantity) if e.available_quantity != None else 0
                e.available_quantity = orig_quantity - int(df[df['name']==name]['quantity'])
                db.session.add(e)
            db.session.commit()
            print()
            today = datetime.datetime.now() 
            dst_path = os.path.join('bills',str(today.year),str(today.month),str(today.day))
            if not os.path.exists(dst_path):
                os.makedirs(dst_path)
            
            doc = docx.Document()
            
            doc.add_heading('Invoice', 0) 
            
 
            d2 = today.strftime("%B %d, %Y")
            doc.add_heading('Name: '+str(b.name).title(), level=2)
            doc.add_heading('Date: '+str(d2), level=2)
            doc.add_heading('',level=1)
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Quantity'
            hdr_cells[2].text = 'Price'
            hdr_cells[3].text = 'Amount'
            records = df[['name','quantity','price','amount']].values
            for name, qty, prc, amt in records:
                row_cells = table.add_row().cells
                row_cells[0].text = str(name)
                row_cells[1].text = str(qty)
                row_cells[2].text = str(prc)
                row_cells[3].text = str(amt)
            
            row_cells = table.add_row().cells
            row_cells[0].text = ''
            row_cells[1].text = ''
            row_cells[2].text = 'Total'
            row_cells[3].text = str(sum(list(df['amount'].values)))
            print(b.name)
            dst_path=os.path.join(basedir,dst_path,'Bill_'+str(today.time()).replace(':','_').replace('.','_')+'.docx')  
            with open(dst_path,'w') as fp:
                pass
            doc.save(dst_path)
            os.system('start '+dst_path)
            return redirect(url_for('billing'))
        return render_template('billing.html',form=form,cb=cb,outs=outs,ub=ub,amount=amount,gb=gb)
    except Exception as e:
        print(e)
        import sys
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(exc_type, fname, exc_tb.tb_lineno)
        return redirect(url_for('billing'))


if __name__=="__main__":
    app.run(debug=True)